from flask import Flask, request, jsonify
from google.cloud import firestore, storage
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.shared import OxmlElement, qn
from datetime import datetime
import tempfile
import os
import re
import traceback
from io import BytesIO
from PIL import Image
import requests
import fnmatch

# Initialize Flask app
app = Flask(__name__)

# Initialize Firestore and Storage clients
db = firestore.Client(database='dawood')
storage_client = storage.Client()
bucket = storage_client.bucket('internship-2025-465209.firebasestorage.app')

def log_with_timestamp(message):
    """Log message with timestamp in Pakistan time"""
    now = datetime.now()
    formatted_time = now.strftime("%Y-%m-%d %H:%M:%S.%f")[:-3] + " PKT"
    print(f"{formatted_time} | {message}")

def sanitize_filename_for_storage(email):
    """Convert email to storage-safe filename"""
    if not email:
        return "unknown"
    return email.replace('@', '_at_').replace('.', '_')

def clean_filename(filename):
    """Clean filename by removing special characters"""
    if not filename:
        return "unknown"
    return re.sub(r'[^a-zA-Z0-9_.-]', '', filename)

def get_student_image_path(student_name):
    """Find correct storage path for student images with comprehensive matching"""
    if not student_name:
        return None
    
    clean_name = student_name.lower().replace(' ', '_')
    name_parts = student_name.split()
    first_name = name_parts[0].lower() if name_parts else ''
    last_name = name_parts[-1].lower() if len(name_parts) > 1 else ''
    
    blobs = list(storage_client.list_blobs(bucket, prefix="students/"))
    
    patterns_to_try = [
        f"*{clean_name}*",
        f"*{first_name}_{last_name}*" if last_name else "",
        f"*{first_name}*",
        f"*{first_name}*{last_name}*" if last_name else "",
        f"*{last_name}*{first_name}*" if last_name else "",
        f"*{first_name}*",
        f"*{re.sub(r'[^a-z]', '', first_name)}*",
        f"*{re.sub(r'[^a-z]', '', clean_name)}*"
    ]
    
    patterns_to_try = [p for p in patterns_to_try if p]
    
    for pattern in patterns_to_try:
        for blob in blobs:
            filename = blob.name.lower().split('/')[-1]
            if fnmatch.fnmatch(filename, pattern):
                log_with_timestamp(f"✅ Found match for {student_name}: {blob.name}")
                return blob.name
    
    log_with_timestamp(f"⚠️ No image found for student: {student_name}")
    return None

def get_project_image(project_data):
    """Find project image with multiple fallback patterns"""
    project_id = project_data.get('id', 'unknown')
    project_title = clean_filename(project_data.get('title', 'unknown'))
    original_title = project_data.get('title', 'unknown')
    
    if project_data.get('imageUrl'):
        try:
            image_path = project_data['imageUrl'].split('/o/')[-1].split('?')[0]
            image_path = image_path.replace('%2F', '/').replace('%20', ' ')
            log_with_timestamp(f"🔄 Trying imageUrl path: {image_path}")
            image_stream = download_image_from_storage(image_path)
            if image_stream:
                return image_stream
            
            decoded_path = requests.utils.unquote(image_path)
            if decoded_path != image_path:
                log_with_timestamp(f"🔄 Trying decoded path: {decoded_path}")
                image_stream = download_image_from_storage(decoded_path)
                if image_stream:
                    return image_stream
        except Exception as e:
            log_with_timestamp(f"⚠️ Error processing image URL: {str(e)}")

    patterns_to_try = [
        f"projects/{project_id}.jpg",
        f"projects/{project_id}.png",
        f"projects/project_{project_id}.jpg",
        f"projects/project_{project_id}.png",
        f"projects/{project_title}.jpg",
        f"projects/{project_title}.png",
        f"projects/{original_title}.jpg",
        f"projects/{original_title}.png",
        f"projects/*{project_title[:10]}*.jpg",
        f"projects/*{project_title[:10]}*.png",
        f"projects/dec cap vale.png",
        f"projects/IOTbasedRealTimeApp.jpg",
        f"projects/Scratch.png"
    ]

    for pattern in patterns_to_try:
        if '*' in pattern:
            blobs = list(storage_client.list_blobs(bucket, prefix="projects/"))
            for blob in blobs:
                if fnmatch.fnmatch(blob.name, pattern):
                    log_with_timestamp(f"🔄 Trying wildcard pattern: {blob.name}")
                    image_stream = download_image_from_storage(blob.name)
                    if image_stream:
                        return image_stream
        else:
            log_with_timestamp(f"🔄 Trying direct path: {pattern}")
            image_stream = download_image_from_storage(pattern)
            if image_stream:
                return image_stream

    log_with_timestamp(f"⚠️ No image found for project {project_id} ({project_title})")
    return download_image_from_storage("projects/default_project.jpg")

def download_image_from_storage(image_path):
    """Download image from Firebase Storage"""
    try:
        if not image_path:
            return None
            
        blob = bucket.blob(image_path)
        if not blob.exists():
            log_with_timestamp(f"⚠️ Image not found: {image_path}")
            return None
        
        image_data = blob.download_as_bytes()
        log_with_timestamp(f"✅ Successfully loaded image: {image_path}")
        return BytesIO(image_data)
    except Exception as e:
        log_with_timestamp(f"❌ Error downloading image {image_path}: {str(e)}")
        return None

def resize_image_for_docx(image_stream, max_width=1.5, max_height=1.5):
    """Resize image to fit in document"""
    try:
        image = Image.open(image_stream)
        width, height = image.size
        aspect_ratio = width / height
        
        if width > height:
            new_width = min(max_width, width / 100)
            new_height = new_width / aspect_ratio
        else:
            new_height = min(max_height, height / 100)
            new_width = new_height * aspect_ratio
        
        if new_width > max_width:
            new_width = max_width
            new_height = max_width / aspect_ratio
        if new_height > max_height:
            new_height = max_height
            new_width = max_height * aspect_ratio
        
        return Inches(new_width), Inches(new_height)
    except Exception as e:
        log_with_timestamp(f"❌ Error resizing image: {str(e)}")
        return Inches(1.5), Inches(1.5)

def format_timestamp(timestamp):
    """Format Firestore timestamp to readable string"""
    if not timestamp:
        return "N/A"
    try:
        if hasattr(timestamp, 'strftime'):  # Already a datetime
            return timestamp.strftime("%B %d, %Y at %I:%M %p")
        elif hasattr(timestamp, 'to_datetime'):  # Firestore timestamp
            return timestamp.to_datetime().strftime("%B %d, %Y at %I:%M %p")
        return "N/A"
    except Exception as e:
        log_with_timestamp(f"⚠️ Error formatting timestamp: {str(e)}")
        return "N/A"

def fetch_supervisor_name(supervisor_id):
    """Fetch supervisor name from Firestore"""
    if not supervisor_id:
        return "N/A"
    try:
        supervisor_doc = db.collection('Teacher').document(supervisor_id).get()
        if supervisor_doc.exists:
            return supervisor_doc.to_dict().get('name', 'N/A')
    except Exception as e:
        log_with_timestamp(f"⚠️ Error fetching supervisor {supervisor_id}: {str(e)}")
    return "N/A"

def add_professional_table(doc, students_data, title="Team Members"):
    """Add a professional-looking table with student information"""
    title_paragraph = doc.add_paragraph()
    title_run = title_paragraph.add_run(title)
    title_run.font.size = Pt(14)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(31, 78, 121)
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    if not students_data:
        doc.add_paragraph("No team members found for this project.")
        return
    
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    for cell in table.rows[0].cells:
        cell.width = Inches(1.5)
    
    header_cells = table.rows[0].cells
    headers = ['S.No.', 'Name', 'Role', 'Email']
    
    for i, header in enumerate(headers):
        header_cells[i].text = header
        for paragraph in header_cells[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(11)
                run.font.color.rgb = RGBColor(255, 255, 255)
        
        shading_elm = OxmlElement("w:shd")
        shading_elm.set(qn("w:fill"), "4472C4")
        header_cells[i]._tc.get_or_add_tcPr().append(shading_elm)
    
    for idx, student in enumerate(students_data, 1):
        row_cells = table.add_row().cells
        row_cells[0].text = str(idx)
        row_cells[1].text = student.get('name', 'N/A')
        row_cells[2].text = student.get('role', 'Student')
        row_cells[3].text = student.get('email', 'N/A')
        
        for cell in row_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10)
                    run.font.name = 'Arial'
    
    doc.add_paragraph()

def add_individual_member_details(doc, students_data):
    """Add individual member details in structured format"""
    if not students_data:
        doc.add_paragraph("No team members found for this project.")
        return
    
    for idx, student in enumerate(students_data, 1):
        member_heading = doc.add_paragraph()
        member_run = member_heading.add_run(f"Member {idx}")
        member_run.font.size = Pt(12)
        member_run.font.bold = True
        member_run.font.color.rgb = RGBColor(68, 114, 196)
        
        details = [
            ("Name", student.get('name', 'N/A')),
            ("Email", student.get('email', 'N/A')),
            ("Role", student.get('role', 'Student'))
        ]
        
        for label, value in details:
            para = doc.add_paragraph()
            label_run = para.add_run(f"{label}: ")
            label_run.font.bold = True
            value_run = para.add_run(value)
            value_run.font.size = Pt(11)
        
        doc.add_paragraph()

def create_professional_docx(projects_data):
    """Create a professional DOCX document with all projects"""
    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)
    
    # Cover page
    title = doc.add_paragraph()
    title_run = title.add_run('PROJECT PORTFOLIO REPORT')
    title_run.font.size = Pt(24)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(31, 78, 121)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph()
    subtitle_run = subtitle.add_run(f'Generated on {datetime.now().strftime("%B %d, %Y")}')
    subtitle_run.font.size = Pt(14)
    subtitle_run.font.italic = True
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    count_para = doc.add_paragraph()
    count_run = count_para.add_run(f'Total Projects: {len(projects_data)}')
    count_run.font.size = Pt(12)
    count_run.font.bold = True
    count_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_page_break()
    
    # Process each project
    for idx, project_data in enumerate(projects_data):
        log_with_timestamp(f"📄 Processing project {idx + 1}/{len(projects_data)}: {project_data.get('title', 'Untitled')}")
        
        # Project title
        project_title = doc.add_paragraph()
        title_run = project_title.add_run(f"Project {idx + 1}: {project_data.get('title', 'Untitled Project')}")
        title_run.font.size = Pt(18)
        title_run.font.bold = True
        title_run.font.color.rgb = RGBColor(31, 78, 121)
        project_title.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        # Project overview
        overview_heading = doc.add_paragraph()
        overview_run = overview_heading.add_run('Project Overview')
        overview_run.font.size = Pt(14)
        overview_run.font.bold = True
        overview_run.font.color.rgb = RGBColor(68, 114, 196)
        
        details = [
            ("Description", project_data.get('description', 'No description available')),
            ("Supervisor", project_data.get('supervisor_name', 'N/A')),
            ("Co-Supervisor", project_data.get('co_supervisor_name', 'N/A'))
        ]
        
        for label, value in details:
            para = doc.add_paragraph()
            label_run = para.add_run(f"{label}: ")
            label_run.font.bold = True
            value_run = para.add_run(value)
            value_run.font.size = Pt(11)
        
        # Timestamps
        timestamps_heading = doc.add_paragraph()
        timestamps_run = timestamps_heading.add_run('Timestamps')
        timestamps_run.font.size = Pt(14)
        timestamps_run.font.bold = True
        timestamps_run.font.color.rgb = RGBColor(68, 114, 196)
        
        created_at = format_timestamp(project_data.get('createdAt'))
        if created_at != "N/A":
            para = doc.add_paragraph()
            label_run = para.add_run("Created At: ")
            label_run.font.bold = True
            value_run = para.add_run(created_at)
            value_run.font.size = Pt(11)
        
        # Project image
        image_stream = get_project_image(project_data)
        
        if image_stream:
            try:
                width, height = resize_image_for_docx(image_stream, max_width=4, max_height=3)
                image_stream.seek(0)
                
                img_para = doc.add_paragraph()
                img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = img_para.add_run()
                run.add_picture(image_stream, width=width, height=height)
                
                caption = doc.add_paragraph()
                caption_run = caption.add_run(f"Figure {idx + 1}: {project_data.get('title', 'Project Image')}")
                caption_run.font.size = Pt(10)
                caption_run.font.italic = True
                caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
            except Exception as e:
                log_with_timestamp(f"❌ Error adding project image: {str(e)}")
        
        # Team members
        students_data = project_data.get('students', [])
        
        # Add table view
        add_professional_table(doc, students_data, "Team Members")
        
        # Add individual member details
        team_members_heading = doc.add_paragraph()
        team_members_run = team_members_heading.add_run('Team Members Details')
        team_members_run.font.size = Pt(14)
        team_members_run.font.bold = True
        team_members_run.font.color.rgb = RGBColor(68, 114, 196)
        
        add_individual_member_details(doc, students_data)
        
        # Team photos
        if students_data:
            photos_title = doc.add_paragraph()
            photos_run = photos_title.add_run('Team Photos')
            photos_run.font.size = Pt(14)
            photos_run.font.bold = True
            photos_run.font.color.rgb = RGBColor(68, 114, 196)
            
            photos_per_row = 3
            current_row_para = None
            
            for i, student in enumerate(students_data):
                if i % photos_per_row == 0:
                    if current_row_para:
                        doc.add_paragraph()
                    current_row_para = doc.add_paragraph()
                    current_row_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                student_name = student.get('name', 'Unknown')
                student_email = student.get('email', '')
                
                # Get the correct image path
                image_path = get_student_image_path(student_name)
                image_stream = download_image_from_storage(image_path) if image_path else None
                
                if image_stream:
                    try:
                        width, height = resize_image_for_docx(image_stream)
                        image_stream.seek(0)
                        run = current_row_para.add_run()
                        run.add_picture(image_stream, width=width, height=height)
                        run.add_text(f" {student_name} ")
                    except Exception as e:
                        log_with_timestamp(f"❌ Error adding photo for {student_name}: {str(e)}")
                        current_row_para.add_run(f"[Photo: {student_name}] ")
                else:
                    current_row_para.add_run(f"[{student_name}] ")
        
        # Only add page break if not the last project
        if idx < len(projects_data) - 1:
            doc.add_page_break()
    
    return doc

def fetch_all_projects(status_filter=None, limit=None):
    """Fetch all projects from Firestore with robust student data handling"""
    try:
        log_with_timestamp(f"🔍 Fetching projects from database")
        
        query = db.collection('projects')
        if status_filter:
            query = query.where('status', '==', status_filter)
        query = query.limit(limit)
        projects = list(query.stream())
        
        projects_data = []
        
        for project_doc in projects:
            try:
                if not project_doc.exists:
                    continue
                    
                data = project_doc.to_dict()
                data['id'] = project_doc.id
                
                # Process students
                students = []
                
                # Approach 1: Students array in project document
                if 'students' in data and isinstance(data['students'], list):
                    for student in data['students']:
                        if isinstance(student, dict):
                            students.append({
                                'name': student.get('name', 'Unknown'),
                                'email': student.get('email', ''),
                                'role': student.get('role', 'Student')
                            })
                
                # Approach 2: Students subcollection
                if not students:
                    try:
                        students_ref = db.collection('projects').document(project_doc.id).collection('students')
                        student_docs = students_ref.stream()
                        for student_doc in student_docs:
                            if student_doc.exists:
                                student_data = student_doc.to_dict()
                                students.append({
                                    'name': student_data.get('name', 'Unknown'),
                                    'email': student_data.get('email', ''),
                                    'role': student_data.get('role', 'Student')
                                })
                    except Exception as e:
                        log_with_timestamp(f"⚠️ Error fetching students subcollection for project {project_doc.id}: {str(e)}")
                
                data['students'] = students
                
                # Fetch supervisor names
                data['supervisor_name'] = fetch_supervisor_name(data.get('supervisorId'))
                data['co_supervisor_name'] = fetch_supervisor_name(data.get('coSupervisorId'))
                
                projects_data.append(data)
                
            except Exception as e:
                log_with_timestamp(f"❌ Error processing project {project_doc.id}: {str(e)}")
                continue
        
        log_with_timestamp(f"✅ Successfully fetched {len(projects_data)} projects")
        return projects_data
    
    except Exception as e:
        log_with_timestamp(f"❌ Error fetching projects: {str(e)}")
        raise e

@app.route('/health', methods=['GET'])
def health_check():
    """Health check endpoint"""
    return jsonify({
        'status': 'healthy',
        'timestamp': datetime.now().isoformat(),
        'service': 'DOCX Generator API',
        'database': 'dawood',
        'bucket': 'internship-2025-465209.firebasestorage.app'
    })

@app.route('/test-firestore', methods=['GET'])
def test_firestore():
    """Test Firestore connection"""
    try:
        docs = db.collection('projects').limit(1).stream()
        return jsonify({
            'success': True,
            'message': 'Firestore connection successful',
            'database': 'dawood',
            'sample_doc_count': len(list(docs))
        })
    except Exception as e:
        return jsonify({
            'success': False,
            'message': f'Firestore connection failed: {str(e)}',
            'database': 'dawood'
        }), 500

@app.route('/', methods=['GET', 'POST'])
def generate_docx():
    """Main endpoint for DOCX generation"""
    log_with_timestamp("🌐 Request received at / endpoint")
    
    try:
        if request.method == 'GET':
            return jsonify({
                'service': 'DOCX Generator API',
                'status': 'running',
                'endpoints': {
                    'POST /': 'Generate DOCX for all projects',
                    'GET /health': 'Health check',
                    'GET /test-firestore': 'Test Firestore connection'
                }
            })
        
        elif request.method == 'POST':
            request_data = request.get_json() or {}
            status_filter = request_data.get('status')
            limit = request_data.get('limit', 100)
            
            projects_data = fetch_all_projects(status_filter=status_filter, limit=limit)
            
            if not projects_data:
                return jsonify({
                    'success': False,
                    'message': 'No projects found with the specified criteria'
                }), 404
            
            doc = create_professional_docx(projects_data)
            
            with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp_file:
                doc.save(tmp_file.name)
                tmp_file_path = tmp_file.name
            
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            filename = f'All_Projects_Report_{timestamp}.docx'
            storage_path = f'documents/{filename}'
            
            blob = bucket.blob(storage_path)
            with open(tmp_file_path, 'rb') as file_data:
                blob.upload_from_file(file_data, content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
            
            blob.make_public()
            download_url = blob.public_url
            os.unlink(tmp_file_path)
            
            return jsonify({
                'success': True,
                'message': f'Successfully generated report for {len(projects_data)} projects',
                'download_url': download_url,
                'filename': filename,
                'projects_count': len(projects_data),
                'generated_at': datetime.now().isoformat(),
                'filters_applied': {
                    'status': status_filter,
                    'limit': limit
                }
            })
        
    except Exception as e:
        log_with_timestamp(f"❌ Error in generate_docx: {str(e)}")
        traceback.print_exc()
        return jsonify({
            'success': False,
            'message': f'Internal server error: {str(e)}'
        }), 500

@app.route('/runtime')
def runtime():
    import os
    return {
        "is_docker": os.path.exists('/.dockerenv'),
        "process": os.popen('ps aux').read()
    }

if __name__ == '__main__':
    app.run(host='0.0.0.0', port=int(os.environ.get('PORT', 8080)), debug=True)
